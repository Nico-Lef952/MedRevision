# -*- coding: utf-8 -*-

import argparse
import asyncio
import json
import os
import re
from pathlib import Path
from datetime import datetime, timezone

import docx
from google import genai
from server import db


def load_env():
    p = Path(".env")
    if not p.exists():
        return

    for line in p.read_text(encoding="utf-8", errors="ignore").splitlines():
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        os.environ.setdefault(k.strip(), v.strip())


def read_annale_file(path):
    path = Path(path)

    if path.suffix.lower() == ".docx":
        document = docx.Document(path)
        parts = []

        for p in document.paragraphs:
            txt = p.text.strip()
            if txt:
                parts.append(txt)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)

        return "\n".join(parts)

    return path.read_text(encoding="utf-8", errors="ignore")


def slugify(text):
    text = text.lower()
    text = re.sub(r"[^\w]+", "-", text)
    return text.strip("-")


def extract_json(text):
    text = (text or "").strip()

    if text.startswith("```"):
        text = re.sub(r"^```json", "", text, flags=re.I).strip()
        text = re.sub(r"^```", "", text).strip()
        text = re.sub(r"```$", "", text).strip()

    start = text.find("{")
    end = text.rfind("}")

    if start == -1 or end == -1:
        raise ValueError("JSON introuvable")

    return json.loads(text[start:end + 1])


def normalize_options(options):
    clean = []

    if not isinstance(options, list):
        return clean

    for opt in options:
        if isinstance(opt, dict):
            txt = opt.get("text") or opt.get("label") or ""
            if txt:
                clean.append({
                    "text": txt,
                    "is_correct": bool(opt.get("is_correct", False))
                })

    return clean


def validate_question(q):
    blob = " ".join([
        str(q.get("question", "")),
        str(q.get("vignette", "")),
        str(q.get("explanation", "")),
        str(q.get("answer", ""))
    ]).lower()

    forbidden_image_markers = [
        "image de",
        "voir image",
        "photo de",
        "schéma",
        "schema",
        "iconographie",
        "si qqn",
        "si quelqu",
        "ci-dessous",
        "cette lésion",
        "cette lesion",
        "structure anatomique 1",
        "structure anatomique 8",
        "structure anatomique 10"
    ]

    if any(marker in blob for marker in forbidden_image_markers):
        return None

    if q.get("needs_image") or q.get("image_required"):
        return None

    qtype = q.get("type", "qrm")

    if qtype not in ["qi", "qcm", "qrm", "qroc", "cas_clinique"]:
        qtype = "qrm"

    question = (q.get("question") or "").strip()
    if not question:
        return None

    answer = (q.get("answer") or "").strip()
    explanation = (q.get("explanation") or "").strip()
    vignette = q.get("vignette", "") or ""
    options = normalize_options(q.get("options", []))

    if qtype in ["qi", "qcm", "qrm"]:
        if len(options) < 2:
            return None
        if not any(o["is_correct"] for o in options):
            return None

    if qtype == "cas_clinique":
        if not vignette:
            return None
        if len(options) < 2:
            return None
        if not any(o["is_correct"] for o in options):
            return None

    if qtype == "qroc":
        if not answer:
            return None
        if len(answer.split()) > 2:
            return None
        options = [{"text": answer, "is_correct": True}]

    return {
        "type": qtype,
        "question": question,
        "vignette": vignette,
        "options": options,
        "answer": answer,
        "explanation": explanation,
        "difficulty": q.get("difficulty", 2),
        "rang": q.get("rang", "A"),
        "concepts": q.get("concepts", []),
        "sub_questions": []
    }


async def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--file", required=True)
    parser.add_argument("--subject", required=True)
    parser.add_argument("--year", required=True)
    parser.add_argument("--title", required=True)
    parser.add_argument("--model", default="gemini-2.5-flash")
    args = parser.parse_args()

    load_env()

    text = read_annale_file(args.file)

    user = await db.users.find_one({"email": "admin@medrevision.com"})
    user_id = str(user["_id"])

    subject = await db.subjects.find_one({
        "user_id": user_id,
        "name": {"$regex": args.subject, "$options": "i"}
    })

    if not subject:
        raise SystemExit("Matière introuvable : " + args.subject)

    subject_id = str(subject["_id"])
    annale_id = f"{subject_id}-{args.year}-{slugify(args.title)}"

    # Évite les doublons si tu réimportes la même annale
    await db.questions.delete_many({
        "user_id": user_id,
        "is_annale": True,
        "annale_id": annale_id
    })

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("Clé Gemini introuvable dans .env")

    client = genai.Client(api_key=api_key)

    prompt = f"""
Tu dois convertir une annale d'examen en questions d'entraînement au format MedRevision.

Matière : {args.subject}
Année : {args.year}
Titre : {args.title}

Consignes :
- Conserve l'esprit exact de l'annale.
- Ne crée pas de questions hors sujet.
- Si une question dépend d'une image, d'un schéma, d'une photo, d'une iconographie ou d'un tableau non présent dans le texte, ignore complètement cette question.
- Si le texte contient une mention comme "image", "photo", "schéma", "si quelqu'un l'a", "voir image", "ci-dessous", sans image exploitable, n'importe pas cette question.
- Pour chaque question importée, fournis la bonne réponse et une correction claire.
- Si l'annale est sous forme QCM/QRM, garde des options.
- Si une question nécessite une réponse longue, transforme-la en QRM ou cas clinique, pas en QROC.
- Les QROC doivent avoir une réponse de 1 ou 2 mots maximum.
- Réponds uniquement en JSON valide.

Annale :
{text}

Format attendu :
{{
  "questions": [
    {{
      "type": "qrm | qcm | qi | qroc | cas_clinique",
      "question": "énoncé",
      "vignette": "",
      "options": [
        {{"text": "option", "is_correct": true}},
        {{"text": "option", "is_correct": false}}
      ],
      "answer": "réponse courte",
      "explanation": "correction",
      "difficulty": 2,
      "rang": "A",
      "concepts": ["notion"]
    }}
  ]
}}
"""

    print("Conversion Gemini en cours...")
    res = client.models.generate_content(model=args.model, contents=prompt)
    data = extract_json(res.text or "")

    docs = []

    for raw in data.get("questions", []):
        q = validate_question(raw)
        if not q:
            continue

        docs.append({
            "user_id": user_id,
            "subject_id": subject_id,
            "course_id": None,
            "is_annale": True,
            "annale_id": annale_id,
            "annale_year": str(args.year),
            "annale_title": args.title,
            "created_at": datetime.now(timezone.utc).isoformat(),
            **q
        })

    if docs:
        await db.questions.insert_many(docs)

    print("Annale importée :", args.title)
    print("Questions insérées :", len(docs))
    print("annale_id :", annale_id)


asyncio.run(main())
