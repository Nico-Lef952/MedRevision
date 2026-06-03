# -*- coding: utf-8 -*-

import argparse
import asyncio
import json
import os
import re
from pathlib import Path
from datetime import datetime, timezone

import docx
from google import genai
from server import db


def load_env():
    p = Path(".env")
    if not p.exists():
        return
    for line in p.read_text(encoding="utf-8", errors="ignore").splitlines():
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        os.environ.setdefault(k.strip(), v.strip())


def slugify(text):
    text = text.lower()
    text = re.sub(r"[^\wÀ-ÿ]+", "-", text)
    return text.strip("-")


def guess_year(filename):
    m = re.search(r"(20\d{2})\s*[-_/ ]\s*(20\d{2})", filename)
    if m:
        return f"{m.group(1)}-{m.group(2)}"
    m = re.search(r"(20\d{2})", filename)
    if m:
        return m.group(1)
    return "année inconnue"


def read_file(path):
    path = Path(path)

    if path.suffix.lower() == ".docx":
        document = docx.Document(path)
        parts = []

        for p in document.paragraphs:
            txt = p.text.strip()
            if txt:
                parts.append(txt)

        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
                line = " | ".join(c for c in cells if c)
                if line:
                    parts.append(line)

        return "\n".join(parts)

    return path.read_text(encoding="utf-8", errors="ignore")


def extract_json(text):
    text = (text or "").strip()

    if text.startswith("```"):
        text = re.sub(r"^```json", "", text, flags=re.I).strip()
        text = re.sub(r"^```", "", text).strip()
        text = re.sub(r"```$", "", text).strip()

    start = text.find("{")
    end = text.rfind("}")

    if start == -1 or end == -1:
        raise ValueError("JSON introuvable")

    return json.loads(text[start:end + 1])


def normalize_options(options):
    clean = []
    for opt in options or []:
        if not isinstance(opt, dict):
            continue

        txt = opt.get("text") or opt.get("label") or opt.get("answer") or ""
        if not txt:
            continue

        clean.append({
            "text": txt.strip(),
            "is_correct": bool(opt.get("is_correct", opt.get("correct", False)))
        })

    return clean


def image_dependent(raw):
    blob = " ".join([
        str(raw.get("question", "")),
        str(raw.get("vignette", "")),
        str(raw.get("answer", "")),
        str(raw.get("explanation", ""))
    ]).lower()

    markers = [
        "image de",
        "voir image",
        "photo de",
        "schéma",
        "schema",
        "iconographie",
        "ci-dessous",
        "cette lésion",
        "cette lesion",
        "structure anatomique 1",
        "structure anatomique 2",
        "structure anatomique 3",
        "structure anatomique 4",
        "structure anatomique 5",
        "structure anatomique 6",
        "structure anatomique 7",
        "structure anatomique 8",
        "structure anatomique 9",
        "structure anatomique 10",
        "si qqn",
        "si quelqu"
    ]

    return any(m in blob for m in markers) or raw.get("needs_image") or raw.get("image_required")


def validate_sub_question(sq):
    if image_dependent(sq):
        return None

    q = (sq.get("question") or "").strip()
    opts = normalize_options(sq.get("options", []))

    if not q or len(opts) < 2 or not any(o["is_correct"] for o in opts):
        return None

    return {
        "type": sq.get("type", "qrm"),
        "question": q,
        "options": opts,
        "answer": (sq.get("answer") or "").strip(),
        "explanation": (sq.get("explanation") or "").strip()
    }


def validate_question(raw):
    if not isinstance(raw, dict):
        return None

    if image_dependent(raw):
        return None

    qtype = raw.get("type", "qrm")
    if qtype not in ["qi", "qcm", "qrm", "qroc", "cas_clinique", "dp"]:
        qtype = "qrm"

    question = (raw.get("question") or "").strip()
    if not question:
        return None

    vignette = raw.get("vignette", "") or ""
    answer = (raw.get("answer") or "").strip()
    explanation = (raw.get("explanation") or "").strip()
    options = normalize_options(raw.get("options", []))

    if qtype in ["qi", "qcm", "qrm"]:
        if len(options) < 2:
            return None
        if not any(o["is_correct"] for o in options):
            return None

    if qtype == "cas_clinique":
        if not vignette:
            return None
        if len(options) < 2:
            return None
        if not any(o["is_correct"] for o in options):
            return None

    if qtype == "qroc":
        if not answer:
            return None
        if len(answer.split()) > 2:
            return None
        options = [{"text": answer, "is_correct": True}]

    sub_questions = []
    if qtype == "dp":
        for sq in raw.get("sub_questions", []) or []:
            valid_sq = validate_sub_question(sq)
            if valid_sq:
                sub_questions.append(valid_sq)

        if len(sub_questions) < 2:
            return None

        options = []

    return {
        "type": qtype,
        "question": question,
        "vignette": vignette,
        "options": options,
        "answer": answer,
        "explanation": explanation,
        "difficulty": raw.get("difficulty", 2),
        "rang": raw.get("rang", "A"),
        "concepts": raw.get("concepts", []),
        "sub_questions": sub_questions
    }


async def get_course_context(user_id, subject_id, max_chars=90000):
    courses = await db.courses.find({
        "user_id": user_id,
        "subject_id": subject_id
    }).sort("title", 1).to_list(500)

    parts = []
    for c in courses:
        title = c.get("title", "")
        content = c.get("content", "") or ""

        if len(content) < 50:
            continue

        parts.append(f"\n\n# COURS : {title}\n{content}")

    context = "\n".join(parts)
    return context[:max_chars]


async def import_one_file(client, args, user_id, subject, file_path):
    subject_id = str(subject["_id"])
    title = file_path.stem.strip()
    year = guess_year(file_path.name)
    annale_id = f"{subject_id}-{year}-{slugify(title)}"

    text = read_file(file_path)
    course_context = await get_course_context(user_id, subject_id)

    await db.questions.delete_many({
        "user_id": user_id,
        "is_annale": True,
        "annale_id": annale_id
    })

    prompt = f"""
Tu dois convertir une annale de dermatologie / revêtement cutané en questions d'entraînement MedRevision.

IMPORTANT :
- Tu dois corriger les questions en te basant prioritairement sur les cours fournis.
- Si une information n'est pas dans les cours, tu peux utiliser le raisonnement médical seulement si la réponse est évidente, mais tu dois le signaler dans l'explication.
- Ne crée pas de questions hors sujet.
- Conserve le plus possible les questions et les items de l'annale.
- Si une question dépend d'une image absente, d'une photo, d'une iconographie ou d'un schéma non présent dans le texte, ignore complètement cette question.
- Si l'annale contient un dossier progressif, crée une question de type "dp" avec des sous_questions.
- Pour les QCM/QRM, conserve les options.
- Les QROC doivent avoir une réponse de 1 ou 2 mots maximum.
- Chaque question doit avoir une correction claire dans "explanation".
- Réponds uniquement en JSON valide.

MATIÈRE :
{args.subject}

TITRE ANNALE :
{title}

ANNÉE :
{year}

COURS DE RÉFÉRENCE :
{course_context}

ANNALE À IMPORTER :
{text}

Format JSON attendu :
{{
  "questions": [
    {{
      "type": "qi | qcm | qrm | qroc | cas_clinique | dp",
      "question": "énoncé",
      "vignette": "",
      "options": [
        {{"text": "option A", "is_correct": true}},
        {{"text": "option B", "is_correct": false}}
      ],
      "answer": "réponse courte",
      "explanation": "correction basée sur les cours",
      "difficulty": 2,
      "rang": "A",
      "concepts": ["notion"],
      "sub_questions": [
        {{
          "type": "qrm",
          "question": "sous-question",
          "options": [
            {{"text": "option A", "is_correct": true}},
            {{"text": "option B", "is_correct": false}}
          ],
          "answer": "réponse",
          "explanation": "correction"
        }}
      ]
    }}
  ]
}}
"""

    print("=" * 100)
    print("Import :", file_path.name)
    print("Annale ID :", annale_id)
    print("=" * 100)

    res = client.models.generate_content(model=args.model, contents=prompt)
    data = extract_json(res.text or "")

    docs = []
    for raw in data.get("questions", []):
        q = validate_question(raw)
        if not q:
            continue

        docs.append({
            "user_id": user_id,
            "subject_id": subject_id,
            "course_id": None,
            "is_annale": True,
            "annale_id": annale_id,
            "annale_year": year,
            "annale_title": title,
            "source_file": str(file_path),
            "created_at": datetime.now(timezone.utc).isoformat(),
            **q
        })

    if docs:
        await db.questions.insert_many(docs)

    print("Questions insérées :", len(docs))
    return len(docs)


async def main():
    parser = argparse.ArgumentParser()
    parser.add_argument("--folder", required=True)
    parser.add_argument("--subject", required=True)
    parser.add_argument("--model", default="gemini-2.5-flash")
    args = parser.parse_args()

    load_env()

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("Clé Gemini introuvable dans .env")

    client = genai.Client(api_key=api_key)

    user = await db.users.find_one({"email": "admin@medrevision.com"})
    user_id = str(user["_id"])

    subject = await db.subjects.find_one({
        "user_id": user_id,
        "name": {"$regex": args.subject, "$options": "i"}
    })

    if not subject:
        raise RuntimeError("Matière introuvable : " + args.subject)

    folder = Path(args.folder)
    files = sorted([
        *folder.glob("*.docx"),
        *folder.glob("*.txt"),
        *folder.glob("*.md")
    ])

    if not files:
        raise RuntimeError("Aucune annale trouvée dans : " + str(folder))

    total = 0
    for file_path in files:
        total += await import_one_file(client, args, user_id, subject, file_path)

    print("\nIMPORT TERMINÉ")
    print("Fichiers :", len(files))
    print("Questions insérées :", total)


asyncio.run(main())
