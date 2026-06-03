# -*- coding: utf-8 -*-

import asyncio
import json
import os
import re
from pathlib import Path
from datetime import datetime, timezone

import docx
from google import genai
from server import db


SUBJECT_NAME = "Revêtement cutané"
MODEL = "gemini-2.5-flash"

QCM_FILE = Path("annales/Revêtement cutané/Sujet 24-25 revêtement cutané.docx")
DP_FILE = Path("annales/Revêtement cutané/dp 24 25.docx")


def load_env():
    p = Path(".env")
    if not p.exists():
        return

    for line in p.read_text(encoding="utf-8", errors="ignore").splitlines():
        if not line or line.startswith("#") or "=" not in line:
            continue
        k, v = line.split("=", 1)
        os.environ.setdefault(k.strip(), v.strip())


def read_docx(path):
    doc = docx.Document(path)
    parts = []

    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt:
            parts.append(txt)

    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip().replace("\n", " / ") for cell in row.cells]
            line = " | ".join(c for c in cells if c)
            if line:
                parts.append(line)

    return "\n".join(parts)


def extract_json(text):
    text = (text or "").strip()

    if text.startswith("```"):
        text = re.sub(r"^```json", "", text, flags=re.I).strip()
        text = re.sub(r"^```", "", text).strip()
        text = re.sub(r"```$", "", text).strip()

    start = text.find("{")
    end = text.rfind("}")

    if start == -1 or end == -1:
        raise ValueError("JSON introuvable")

    return json.loads(text[start:end + 1])


def normalize_options(options):
    clean = []

    for opt in options or []:
        if not isinstance(opt, dict):
            continue

        text = opt.get("text") or opt.get("label") or ""
        if not text:
            continue

        clean.append({
            "text": text.strip(),
            "is_correct": bool(opt.get("is_correct", opt.get("correct", False)))
        })

    return clean


def validate_qcm(q):
    qtype = q.get("type", "qrm")
    if qtype not in ["qi", "qcm", "qrm"]:
        qtype = "qrm"

    question = (q.get("question") or "").strip()
    options = normalize_options(q.get("options", []))
    explanation = (q.get("explanation") or "").strip()
    answer = (q.get("answer") or "").strip()

    if not question:
        return None

    if len(options) < 2:
        return None

    if not any(o["is_correct"] for o in options):
        return None

    if not explanation:
        explanation = "Correction à vérifier."

    return {
        "type": qtype,
        "question": question,
        "vignette": q.get("vignette", "") or "",
        "options": options,
        "answer": answer,
        "explanation": explanation,
        "difficulty": q.get("difficulty", 2),
        "rang": q.get("rang", "A"),
        "concepts": q.get("concepts", []),
        "sub_questions": []
    }


def validate_dp(dp, expected_count):
    question = (dp.get("question") or "").strip()
    vignette = (dp.get("vignette") or "").strip()
    sub_questions = []

    if not question:
        question = "Dossier progressif"

    if not vignette:
        raise ValueError("DP sans vignette")

    for sq in dp.get("sub_questions", []) or []:
        q = (sq.get("question") or "").strip()
        opts = normalize_options(sq.get("options", []))
        explanation = (sq.get("explanation") or "").strip()
        answer = (sq.get("answer") or "").strip()

        if not q:
            continue
        if len(opts) < 2:
            continue
        if not any(o["is_correct"] for o in opts):
            continue

        sub_questions.append({
            "type": sq.get("type", "qrm"),
            "question": q,
            "options": opts,
            "answer": answer,
            "explanation": explanation or "Correction à vérifier."
        })

    if len(sub_questions) != expected_count:
        raise ValueError(f"DP invalide : attendu {expected_count} sous-questions, obtenu {len(sub_questions)}")

    return {
        "type": "dp",
        "question": question,
        "vignette": vignette,
        "options": [],
        "answer": dp.get("answer", "") or "",
        "explanation": dp.get("explanation", "") or "",
        "difficulty": dp.get("difficulty", 2),
        "rang": dp.get("rang", "A"),
        "concepts": dp.get("concepts", []),
        "sub_questions": sub_questions
    }


async def get_course_context(user_id, subject_id):
    courses = await db.courses.find({
        "user_id": user_id,
        "subject_id": subject_id
    }).sort("title", 1).to_list(500)

    parts = []
    for c in courses:
        title = c.get("title", "")
        content = c.get("content", "") or ""
        if len(content) > 100:
            parts.append(f"\n\n# {title}\n{content}")

    return "\n".join(parts)[:90000]


async def main():
    load_env()

    api_key = os.environ.get("GEMINI_API_KEY") or os.environ.get("GOOGLE_API_KEY")
    if not api_key:
        raise RuntimeError("Clé Gemini introuvable")

    client = genai.Client(api_key=api_key)

    user = await db.users.find_one({"email": "admin@medrevision.com"})
    user_id = str(user["_id"])

    subject = await db.subjects.find_one({
        "user_id": user_id,
        "name": {"$regex": SUBJECT_NAME, "$options": "i"}
    })

    if not subject:
        raise RuntimeError("Matière Revêtement cutané introuvable")

    subject_id = str(subject["_id"])
    course_context = await get_course_context(user_id, subject_id)

    # Nettoyage ancien import dermato
    deleted = await db.questions.delete_many({
        "user_id": user_id,
        "is_annale": True,
        "subject_id": subject_id
    })
    print("Anciennes annales dermato supprimées :", deleted.deleted_count)

    # =========================
    # 1. QCM 30 questions
    # =========================
    qcm_text = read_docx(QCM_FILE)
    qcm_annale_id = f"{subject_id}-2024-2025-sujet-revetement-cutane"

    qcm_prompt = f"""
Tu dois convertir cette annale de Revêtement cutané en EXACTEMENT 30 questions MedRevision.

Règles :
- Tu dois conserver les 30 questions de l'annale.
- Ne résume pas.
- Ne fusionne pas les questions.
- Chaque question doit garder ses options.
- Corrige les réponses avec Gemini en t'appuyant d'abord sur les cours fournis.
- Chaque question doit avoir une explanation claire.
- Si une question dépend d'une image absente, garde-la seulement si le texte permet de répondre sans image.
- Réponds uniquement en JSON valide.

COURS :
{course_context}

ANNALE QCM :
{qcm_text}

Format :
{{
  "questions": [
    {{
      "type": "qrm",
      "question": "énoncé",
      "vignette": "",
      "options": [
        {{"text": "A. option", "is_correct": true}},
        {{"text": "B. option", "is_correct": false}}
      ],
      "answer": "réponse courte",
      "explanation": "correction",
      "difficulty": 2,
      "rang": "A",
      "concepts": ["notion"]
    }}
  ]
}}
"""

    print("\nImport QCM 30 questions...")
    res = client.models.generate_content(model=MODEL, contents=qcm_prompt)
    data = extract_json(res.text or "")

    qcm_docs = []
    for raw in data.get("questions", []):
        q = validate_qcm(raw)
        if q:
            qcm_docs.append({
                "user_id": user_id,
                "subject_id": subject_id,
                "course_id": None,
                "is_annale": True,
                "annale_id": qcm_annale_id,
                "annale_year": "2024-2025",
                "annale_title": "Sujet 24-25 revêtement cutané",
                "source_file": str(QCM_FILE),
                "created_at": datetime.now(timezone.utc).isoformat(),
                **q
            })

    if len(qcm_docs) != 30:
        Path("debug_dermato_qcm_response.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        raise RuntimeError(f"QCM import refusé : attendu 30 questions, obtenu {len(qcm_docs)}. Voir debug_dermato_qcm_response.json")

    await db.questions.insert_many(qcm_docs)
    print("QCM importé : 30 questions")

    # =========================
    # 2. DP : 2 dossiers, 7 et 5 sous-questions
    # =========================
    dp_text = read_docx(DP_FILE)
    dp_annale_id = f"{subject_id}-2024-2025-dp-revetement-cutane"

    dp_prompt = f"""
Tu dois convertir ce fichier de DP de Revêtement cutané en EXACTEMENT 2 dossiers progressifs MedRevision.

Règles strictes :
- Il y a exactement 2 DP.
- Le premier DP doit contenir exactement 7 sous-questions.
- Le deuxième DP doit contenir exactement 5 sous-questions.
- Chaque DP doit avoir une vignette complète.
- Chaque sous-question doit avoir un énoncé compréhensible avec la vignette.
- Ne supprime aucune sous-question.
- Ne résume pas.
- Conserve les options.
- Corrige avec Gemini en t'appuyant d'abord sur les cours fournis.
- Chaque sous-question doit avoir une explanation claire.
- Réponds uniquement en JSON valide.

COURS :
{course_context}

FICHIER DP :
{dp_text}

Format :
{{
  "dps": [
    {{
      "type": "dp",
      "question": "titre du DP",
      "vignette": "vignette clinique complète",
      "explanation": "correction globale courte",
      "difficulty": 2,
      "rang": "A",
      "concepts": ["notion"],
      "sub_questions": [
        {{
          "type": "qrm",
          "question": "énoncé sous-question",
          "options": [
            {{"text": "A. option", "is_correct": true}},
            {{"text": "B. option", "is_correct": false}}
          ],
          "answer": "réponse courte",
          "explanation": "correction"
        }}
      ]
    }}
  ]
}}
"""

    print("\nImport DP 7 + 5...")
    res = client.models.generate_content(model=MODEL, contents=dp_prompt)
    data = extract_json(res.text or "")

    Path("debug_dermato_dp_response.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")

    dps_raw = data.get("dps", [])
    print("Debug DP sauvegardé dans debug_dermato_dp_response.json")
    for idx_dbg, dp_dbg in enumerate(dps_raw, 1):
        print(f"DP {idx_dbg} brut :", len(dp_dbg.get("sub_questions", []) or []), "sous-questions")

    if len(dps_raw) != 2:
        Path("debug_dermato_dp_response.json").write_text(json.dumps(data, ensure_ascii=False, indent=2), encoding="utf-8")
        raise RuntimeError(f"DP import refusé : attendu 2 DP, obtenu {len(dps_raw)}. Voir debug_dermato_dp_response.json")

    expected = [7, 5]
    dp_docs = []

    for idx, raw_dp in enumerate(dps_raw):
        dp = validate_dp(raw_dp, expected[idx])
        dp_docs.append({
            "user_id": user_id,
            "subject_id": subject_id,
            "course_id": None,
            "is_annale": True,
            "annale_id": dp_annale_id,
            "annale_year": "2024-2025",
            "annale_title": "DP 24-25 revêtement cutané",
            "source_file": str(DP_FILE),
            "source_key": f"dp_dermato_2024_2025_{idx+1}",
            "created_at": datetime.now(timezone.utc).isoformat(),
            **dp
        })

    await db.questions.insert_many(dp_docs)
    print("DP importés : 2 dossiers, 7 + 5 sous-questions")

    print("\nIMPORT DERMATO OK")
    print("Sujet QCM : 30 questions")
    print("DP : 2 dossiers, 12 sous-questions au total")


asyncio.run(main())
